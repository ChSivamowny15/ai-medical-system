from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

# =====================================
# REPORT FOLDER
# =====================================

BASE_DIR = os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))
)

REPORT_FOLDER = os.path.join(
    BASE_DIR,
    "reports"
)

os.makedirs(REPORT_FOLDER, exist_ok=True)

# =====================================
# GENERATE REPORT
# =====================================

def generate_report(
    prediction,
    confidence,
    scan_type,
    patient_name,
    patient_age,
    patient_gender,
    patient_symptoms,
    image_path
):

    # =====================================
    # FIX VALUES
    # =====================================

    prediction = str(prediction).strip().lower()

    scan_type = str(scan_type).strip().lower()

    # =====================================
    # DOCUMENT
    # =====================================

    document = Document()

    # =====================================
    # TITLE
    # =====================================

    title = document.add_heading(
        "AI MEDICAL DIAGNOSTIC REPORT",
        level=1
    )

    title.runs[0].font.size = Pt(22)

    # =====================================
    # BASIC DETAILS
    # =====================================

    document.add_paragraph(
        f"Prediction : {prediction.upper()}"
    )

    document.add_paragraph(
        f"Confidence : {confidence}%"
    )

    document.add_paragraph(
        f"Scan Type : {scan_type.upper()}"
    )

    document.add_paragraph(
        f"Generated Time : {datetime.now()}"
    )

    # =====================================
    # PATIENT DETAILS
    # =====================================

    document.add_heading(
        "Patient Details",
        level=2
    )

    document.add_paragraph(
        f"Patient Name : {patient_name}"
    )

    document.add_paragraph(
        f"Age : {patient_age}"
    )

    document.add_paragraph(
        f"Gender : {patient_gender}"
    )

    document.add_paragraph(
        f"Symptoms : {patient_symptoms}"
    )

    # =====================================
    # MRI REPORT
    # =====================================

    if scan_type == "mri":

        document.add_heading(
            "MRI Findings",
            level=2
        )

        if prediction == "yes":

            document.add_paragraph(
                "The MRI scan indicates possible abnormal tissue growth in the brain region which may represent a tumor or neurological abnormality."
            )

            document.add_heading(
                "Possible Symptoms",
                level=2
            )

            document.add_paragraph(
                "- Severe headaches\n"
                "- Vomiting\n"
                "- Blurred vision\n"
                "- Memory loss\n"
                "- Difficulty in balance"
            )

            document.add_heading(
                "Recommended Medical Actions",
                level=2
            )

            document.add_paragraph(
                "- Immediate neurologist consultation advised\n"
                "- MRI with contrast recommended\n"
                "- Continuous monitoring required"
            )

            document.add_heading(
                "Possible Treatment",
                level=2
            )

            document.add_paragraph(
                "- Surgery\n"
                "- Radiation therapy\n"
                "- Chemotherapy\n"
                "- Targeted therapy"
            )

            document.add_heading(
                "Precautions",
                level=2
            )

            document.add_paragraph(
                "- Avoid stress and heavy workload\n"
                "- Maintain healthy sleep cycle\n"
                "- Follow medication schedule properly"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "HIGH RISK - Immediate medical attention required."
            )

        else:

            document.add_paragraph(
                "The MRI scan appears normal with no major abnormalities detected."
            )

            document.add_heading(
                "Health Status",
                level=2
            )

            document.add_paragraph(
                "- Brain appears healthy\n"
                "- No tumor detected\n"
                "- Neurological structure normal"
            )

            document.add_heading(
                "Recommendations",
                level=2
            )

            document.add_paragraph(
                "- Continue healthy lifestyle\n"
                "- Exercise regularly\n"
                "- Maintain regular medical checkups"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "LOW RISK - No major abnormality detected."
            )

    # =====================================
    # CT REPORT
    # =====================================

    elif scan_type == "ct":

        document.add_heading(
            "CT Scan Findings",
            level=2
        )

        if prediction == "yes":

            document.add_paragraph(
                "The CT scan indicates possible lung abnormalities or infection patterns in the respiratory region."
            )

            document.add_heading(
                "Possible Symptoms",
                level=2
            )

            document.add_paragraph(
                "- Chest pain\n"
                "- Difficulty in breathing\n"
                "- Persistent cough\n"
                "- Fever\n"
                "- Fatigue"
            )

            document.add_heading(
                "Recommended Medical Actions",
                level=2
            )

            document.add_paragraph(
                "- Consult pulmonologist immediately\n"
                "- Additional CT imaging recommended\n"
                "- Respiratory examination advised"
            )

            document.add_heading(
                "Possible Treatment",
                level=2
            )

            document.add_paragraph(
                "- Respiratory therapy\n"
                "- Antibiotics\n"
                "- Oxygen support\n"
                "- Lung rehabilitation"
            )

            document.add_heading(
                "Precautions",
                level=2
            )

            document.add_paragraph(
                "- Avoid smoking\n"
                "- Stay away from pollution\n"
                "- Practice breathing exercises"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "MEDIUM TO HIGH RISK - Respiratory monitoring recommended."
            )

        else:

            document.add_paragraph(
                "The CT scan appears normal with no major lung abnormalities detected."
            )

            document.add_heading(
                "Health Status",
                level=2
            )

            document.add_paragraph(
                "- Lungs appear healthy\n"
                "- No infection patterns found\n"
                "- Respiratory system stable"
            )

            document.add_heading(
                "Recommendations",
                level=2
            )

            document.add_paragraph(
                "- Maintain healthy breathing habits\n"
                "- Regular exercise recommended\n"
                "- Avoid polluted environments"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "LOW RISK - No major abnormality detected."
            )

    # =====================================
    # UV REPORT
    # =====================================

    elif scan_type == "uv":

        document.add_heading(
            "UV Scan Findings",
            level=2
        )

        if prediction == "yes":

            document.add_paragraph(
                "The UV scan detected possible skin lesion or abnormal pigmentation patterns."
            )

            document.add_heading(
                "Possible Symptoms",
                level=2
            )

            document.add_paragraph(
                "- Skin irritation\n"
                "- Itching\n"
                "- Dark patches\n"
                "- Burning sensation\n"
                "- Skin redness"
            )

            document.add_heading(
                "Recommended Medical Actions",
                level=2
            )

            document.add_paragraph(
                "- Dermatologist consultation advised\n"
                "- Skin biopsy may be required\n"
                "- Advanced skin examination recommended"
            )

            document.add_heading(
                "Possible Treatment",
                level=2
            )

            document.add_paragraph(
                "- Skin creams\n"
                "- Laser therapy\n"
                "- Anti-allergic medications\n"
                "- Dermatological therapy"
            )

            document.add_heading(
                "Precautions",
                level=2
            )

            document.add_paragraph(
                "- Avoid direct sunlight exposure\n"
                "- Use sunscreen regularly\n"
                "- Maintain proper skin hygiene"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "MEDIUM RISK - Dermatological examination recommended."
            )

        else:

            document.add_paragraph(
                "The UV scan appears normal with no major skin abnormalities detected."
            )

            document.add_heading(
                "Health Status",
                level=2
            )

            document.add_paragraph(
                "- Skin condition appears healthy\n"
                "- No lesion patterns found\n"
                "- No infection signs observed"
            )

            document.add_heading(
                "Recommendations",
                level=2
            )

            document.add_paragraph(
                "- Continue proper skincare\n"
                "- Maintain skin hygiene\n"
                "- Use sunscreen regularly"
            )

            document.add_heading(
                "Risk Level",
                level=2
            )

            document.add_paragraph(
                "LOW RISK - Skin condition appears normal."
            )

    # =====================================
    # DISCLAIMER
    # =====================================

    document.add_heading(
        "Disclaimer",
        level=2
    )

    document.add_paragraph(
        "This AI-generated report should not be considered a final medical diagnosis. Please consult a certified medical professional for further evaluation."
    )

    # =====================================
    # DOCTOR SIGNATURE
    # =====================================

    document.add_heading(
        "Authorized By",
        level=2
    )

    document.add_paragraph(
        "AI Medical Diagnostic System\n"
        "Verified Healthcare Assistance Platform"
    )

    # =====================================
    # SAVE REPORT
    # =====================================

    filename = f"{scan_type}_report.docx"

    report_path = os.path.join(
        REPORT_FOLDER,
        filename
    )

    document.save(report_path)

    return report_path